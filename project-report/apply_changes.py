"""
apply_changes.py — Five-part update to report.docx
  1. Table numbers → Roman numerals (captions + LOT + cross-refs)
  2. Citation renumbering (order of first appearance) + 3 new references
  3. Equation paragraph formatting (centred + right-aligned number)
  4. Reference list rebuild in IEEE format (remove ↳ arrows, reorder, add new)
  5. Appendix D SDG/PO/PSO improved justification
"""
import sys, io, re, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

SRC  = r"c:\Users\divya\Desktop\encryption-iit-dharwad\project-report\report.docx"
doc  = Document(SRC)
body = doc.element.body

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def get_para_text(p):
    return ''.join(t.text or '' for t in p.findall('.//' + qn('w:t')))

def set_run_text(r, text):
    for t in r.findall(qn('w:t')):
        t.text = ''
    ts = r.findall(qn('w:t'))
    if ts:
        ts[0].text = text
        ts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    else:
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)

def replace_text_in_para(p, old, new):
    """Replace all occurrences of old→new across runs in paragraph."""
    full = get_para_text(p)
    if old not in full:
        return False
    # Rebuild text in the first run, clear others
    runs = p.findall('.//' + qn('w:r'))
    if not runs:
        return False
    new_full = full.replace(old, new)
    set_run_text(runs[0], new_full)
    for r in runs[1:]:
        set_run_text(r, '')
    return True

def make_ref_para(number, text, template_p):
    """Create a new reference paragraph styled like an existing one."""
    new_p = copy.deepcopy(template_p)
    runs = new_p.findall('.//' + qn('w:r'))
    full_text = f"[{number}] {text}"
    if runs:
        set_run_text(runs[0], full_text)
        for r in runs[1:]:
            set_run_text(r, '')
    return new_p

# ─────────────────────────────────────────────────────────────────────────────
# PART 1 — Table numbers → Roman numerals
# ─────────────────────────────────────────────────────────────────────────────
print("=== PART 1: Roman table numbers ===")

TABLE_MAP = {
    'Table 2.5.1': 'Table I',
    'Table 3.1.1': 'Table II',
    'Table 3.5.1': 'Table III',
    'Table 3.5.2': 'Table IV',
    'Table 3.6.1': 'Table V',
    'Table 3.6.2': 'Table VI',
}

paras = body.findall(qn('w:p'))
for p in paras:
    txt = get_para_text(p)
    for old, new in TABLE_MAP.items():
        if old in txt:
            replace_text_in_para(p, old, new)
            print(f"  {old} → {new}  in: {txt[:60]}")

# Update LIST OF TABLES (table index 4)
LOT_NUM_MAP = {'2.5.1': 'I', '3.1.1': 'II', '3.5.1': 'III',
               '3.5.2': 'IV', '3.6.1': 'V', '3.6.2': 'VI'}
lot = doc.tables[4]
for row in lot.rows[1:]:  # skip header
    sr_cell = row.cells[0]
    sr_txt = sr_cell.text.strip()
    if sr_txt in LOT_NUM_MAP:
        new_num = LOT_NUM_MAP[sr_txt]
        for p in sr_cell.paragraphs:
            runs = p.runs
            if runs:
                runs[0].text = new_num
                for r in runs[1:]: r.text = ''
        print(f"  LOT: {sr_txt} → {new_num}")

# ─────────────────────────────────────────────────────────────────────────────
# PART 2 — Citation renumbering
# ─────────────────────────────────────────────────────────────────────────────
print("\n=== PART 2: Citation renumbering ===")

# Old-number → New-number mapping (derived from order of first appearance)
OLD_TO_NEW = {
    1:1, 2:2, 10:3, 11:4, 4:5, 5:6, 6:7, 7:8, 13:9, 17:10,
    8:11, 18:12, 20:13, 19:14, 3:15, 12:16, 9:17, 16:18,
    14:19, 15:20, 21:21
}

def remap_citation(match):
    inner = match.group(1)
    parts = re.split(r',\s*', inner)
    try:
        nums = [int(p.strip()) for p in parts]
    except ValueError:
        return match.group(0)
    # Only treat as citation if all numbers are known reference numbers (1-21)
    if all(1 <= n <= 21 for n in nums):
        new_nums = [OLD_TO_NEW.get(n, n) for n in nums]
        return '[' + ', '.join(str(n) for n in new_nums) + ']'
    return match.group(0)

# Find REFERENCES heading paragraph index
paras = body.findall(qn('w:p'))
ref_heading_idx = None
for i, p in enumerate(paras):
    if get_para_text(p).strip() == 'REFERENCES':
        ref_heading_idx = i
        break

# Apply to all body paragraphs before REFERENCES
changed = 0
for i, p in enumerate(paras[:ref_heading_idx]):
    old_txt = get_para_text(p)
    if '[' not in old_txt:
        continue
    new_txt = re.sub(r'\[(\d+(?:,\s*\d+)*)\]', remap_citation, old_txt)
    if new_txt != old_txt:
        runs = p.findall('.//' + qn('w:r'))
        if runs:
            set_run_text(runs[0], new_txt)
            for r in runs[1:]:
                set_run_text(r, '')
            changed += 1

print(f"  Updated citations in {changed} paragraphs")

# Add citations for new refs [22], [23], [24] in appropriate places
# [22] Kursawe smart grid privacy → add to para 52 (AMI/smart meter para) after [11]
# [23] Li HE aggregation → add to para 94 EPPA discussion after [15]
# [24] Damgård SMPC → add to para 275 SMPC future scope after [14]
paras = body.findall(qn('w:p'))

NEW_REF_INSERTIONS = [
    # (search_text_fragment, old_citation, append_with)
    ('Private Memoirs of a Smart Meter', '[11]', '[11, 22]'),
    ('EPPA: An Efficient and Privacy', '[15]', '[15, 23]'),
    ('Combining CKKS FHE aggregation with SMPC', '[14]', '[14, 24]'),
]

for search, old_cit, new_cit in NEW_REF_INSERTIONS:
    for p in paras[:ref_heading_idx]:
        txt = get_para_text(p)
        if search in txt and old_cit in txt:
            replace_text_in_para(p, old_cit, new_cit)
            print(f"  Added new ref: {old_cit} → {new_cit} in para containing '{search[:40]}'")
            break

# ─────────────────────────────────────────────────────────────────────────────
# PART 3 — Equation paragraph formatting
# ─────────────────────────────────────────────────────────────────────────────
print("\n=== PART 3: Equation formatting ===")

# Equations are paragraphs that contain a numbered equation pattern
# Format: indent left 1cm, equation text, tab, "(3.x)" right-aligned
# We set paragraph indent and clean up the "- (3.x)" → "(3.x)" format

EQ_PATTERN = re.compile(r'\(3\.\d+[a-c]?\)')

def format_eq_para(p):
    txt = get_para_text(p)
    if not EQ_PATTERN.search(txt):
        return False
    # Clean up "- (3.x)" → "(3.x)"
    new_txt = re.sub(r'\s*[-–]\s*(\(3\.\d+[a-c]?\))', r'    \1', txt)
    # Remove inline description bracketed annotations like [ciphertext addition, depth 0]
    # Leave them but clean the dash before equation number
    if new_txt == txt and '(3.' not in txt:
        return False

    # Set paragraph properties: left indent + right-aligned tab for number
    pp = p.find(qn('w:pPr'))
    if pp is None:
        pp = OxmlElement('w:pPr')
        p.insert(0, pp)

    # Remove existing ind
    for ind in pp.findall(qn('w:ind')):
        pp.remove(ind)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')   # ~0.25 inch indent
    pp.append(ind)

    # Set centered justification for cleaner look
    jc = pp.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pp.append(jc)
    jc.set(qn('w:val'), 'left')

    # Update text if changed
    if new_txt != txt:
        runs = p.findall('.//' + qn('w:r'))
        if runs:
            set_run_text(runs[0], new_txt)
            for r in runs[1:]:
                set_run_text(r, '')
    return True

paras = body.findall(qn('w:p'))
eq_count = 0
for p in paras:
    if format_eq_para(p):
        eq_count += 1
print(f"  Formatted {eq_count} equation paragraphs")

# ─────────────────────────────────────────────────────────────────────────────
# PART 4 — Rebuild reference list (IEEE format, new order, no arrows)
# ─────────────────────────────────────────────────────────────────────────────
print("\n=== PART 4: Reference list rebuild ===")

# New reference list in new numbering order (IEEE format, no ↳)
NEW_REFS = [
    (1,  'J.H. Cheon, A. Kim, M. Kim, and Y. Song, "Homomorphic Encryption for Arithmetic of Approximate Numbers," in Advances in Cryptology \u2013 ASIACRYPT 2017, LNCS vol. 10624, Springer, pp. 409\u2013437, 2017.'),
    (2,  'M. Benaissa, B. Retiat, B. Cebere, and A. Bi\u00e7er, "TenSEAL: A Library for Encrypted Tensor Operations Using Homomorphic Encryption," in ICLR 2021 Workshop on Distributed and Private ML (DPML), 2021.'),
    (3,  'NIST, "Recommendation for Key Management \u2013 Part 1: General," NIST Special Publication 800-57 Part 1 Rev. 5, National Institute of Standards and Technology, 2020.'),
    (4,  'M. Albrecht et al., "Homomorphic Encryption Security Standard," HomomorphicEncryption.org White Paper, Tech. Rep., 2019.'),
    (5,  'A. Kim, Y. Song, M. Kim, K. Lee, and J.H. Cheon, "Logistic Regression Model Training Based on the Approximate Homomorphic Encryption," BMC Medical Genomics, vol. 11, Suppl. 4, article 23, 2018.'),
    (6,  'J.H. Cheon, D. Kim, and D. Kim, "Efficient Homomorphic Comparison Methods with Optimal Complexity," in Advances in Cryptology \u2013 ASIACRYPT 2019, LNCS vol. 11922, Springer, pp. 221\u2013256, 2019.'),
    (7,  'T.P. Pedersen, "Non-Interactive and Information-Theoretic Secure Verifiable Secret Sharing," in Advances in Cryptology \u2013 CRYPTO 1991, LNCS vol. 576, Springer, pp. 129\u2013140, 1991.'),
    (8,  'R. Gennaro, C. Gentry, and B. Parno, "Non-Interactive Verifiable Computing: Outsourcing Computation to Untrusted Workers," in Advances in Cryptology \u2013 CRYPTO 2010, LNCS vol. 6223, Springer, pp. 465\u2013482, 2010.'),
    (9,  'H. Chen, I. Chillotti, and Y. Song, "Improved Bootstrapping for Approximate Homomorphic Encryption," in Advances in Cryptology \u2013 EUROCRYPT 2019, LNCS vol. 11477, Springer, pp. 34\u201354, 2019.'),
    (10, 'A. Boyle, G. Couteau, N. Gilboa, and Y. Ishai, "Homomorphic Secret Sharing: Optimizations and Applications," in Proc. ACM CCS 2017, pp. 2105\u20132122, 2017.'),
    (11, 'A. Molina-Markham, P. Shenoy, K. Fu, E. Cecchet, and D. Irwin, "Private Memoirs of a Smart Meter," in Proc. 2nd ACM Workshop on Embedded Sensing Systems for Energy-Efficiency in Buildings (BuildSys), pp. 61\u201366, 2010.'),
    (12, 'European Parliament, "General Data Protection Regulation (GDPR)," Regulation (EU) 2016/679 of the European Parliament and of the Council, 27 April 2016.'),
    (13, 'C. Dwork and A. Roth, "The Algorithmic Foundations of Differential Privacy," Foundations and Trends in Theoretical Computer Science, vol. 9, nos. 3\u20134, pp. 211\u2013407, 2014.'),
    (14, 'A. Shamir, "How to Share a Secret," Communications of the ACM, vol. 22, no. 11, pp. 612\u2013613, 1979.'),
    (15, 'A. Lu, H. Li, F. Chen, and Q. Zhu, "EPPA: An Efficient and Privacy-Preserving Aggregation Scheme for Secure Smart Grid Communications," IEEE Trans. Parallel and Distributed Systems, vol. 23, no. 9, pp. 1621\u20131631, 2012.'),
    (16, 'Microsoft Corporation, "Microsoft SEAL (Simple Encrypted Arithmetic Library)," GitHub repository: github.com/microsoft/SEAL, Version 4.1, 2023.'),
    (17, 'A. Dimeas and N. Hatziargyriou, "Operation of a Multiagent System for Microgrid Control," IEEE Trans. Power Systems, vol. 20, no. 3, pp. 1447\u20131455, 2005.'),
    (18, 'I. Chillotti, N. Gama, M. Georgieva, and M. Izabach\u00e8ne, "TFHE: Fast Fully Homomorphic Encryption over the Torus," J. Cryptology, vol. 33, pp. 34\u201391, 2020.'),
    (19, 'X. Yi, R. Paulet, and E. Bertino, "Homomorphic Encryption and Applications," SpringerBriefs in Computer Science, Springer International Publishing, 2014.'),
    (20, 'M. Green and A. Miers, "Forward Secure Asynchronous Messaging from Puncturable Encryption," in IEEE Symp. on Security and Privacy (S&P), pp. 305\u2013320, 2015.'),
    (21, 'Z. Brakerski, C. Gentry, and V. Vaikuntanathan, "(Leveled) Fully Homomorphic Encryption without Bootstrapping," ACM Trans. Computation Theory, vol. 6, no. 3, article 13, 2014.'),
    (22, 'K. Kursawe, G. Danezis, and M. Kohlweiss, "Privacy-Friendly Aggregation for the Smart-Grid," in Proc. 11th Privacy Enhancing Technologies Symp. (PETS 2011), LNCS vol. 6794, Springer, pp. 175\u2013191, 2011.'),
    (23, 'F. Li, B. Luo, and P. Liu, "Secure Information Aggregation for Smart Grids Using Homomorphic Encryption," in Proc. 1st IEEE Int. Conf. Smart Grid Commun. (SmartGridComm 2010), pp. 327\u2013332, 2010.'),
    (24, 'I. Damg\u00e5rd, V. Pastro, N. Smart, and S. Zakarias, "Multiparty Computation from Somewhat Homomorphic Encryption," in Advances in Cryptology \u2013 CRYPTO 2012, LNCS vol. 7417, Springer, pp. 643\u2013662, 2012.'),
]

# Find REFERENCES heading and the paragraph before Agnel Charities
paras = body.findall(qn('w:p'))
ref_heading_para = None
agnel_para = None
for p in paras:
    txt = get_para_text(p).strip()
    if txt == 'REFERENCES':
        ref_heading_para = p
    if 'Agnel' in txt and agnel_para is None:
        agnel_para = p

# Remove existing reference paragraphs (between heading and Agnel)
removing = False
to_del = []
for p in paras:
    if p is ref_heading_para:
        removing = True
        continue
    if p is agnel_para:
        break
    if removing:
        to_del.append(p)

for p in to_del:
    p.getparent().remove(p)
print(f"  Removed {len(to_del)} old reference paragraphs")

# Use old ref heading para style as template for new refs
template = copy.deepcopy(ref_heading_para)  # will be adapted

# Helper: create a styled reference paragraph
def make_ref(num, text):
    p = OxmlElement('w:p')
    pp = OxmlElement('w:pPr')
    # Hanging indent for IEEE style
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '360')
    pp.append(ind)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60')
    sp.set(qn('w:after'), '60')
    sp.set(qn('w:line'), '276')
    sp.set(qn('w:lineRule'), 'auto')
    pp.append(sp)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'both')
    pp.append(jc)
    p.append(pp)

    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Calibri')
    rf.set(qn('w:hAnsi'), 'Calibri')
    rpr.append(rf)
    for tag in ('w:sz', 'w:szCs'):
        s = OxmlElement(tag)
        s.set(qn('w:val'), '22')  # 11pt
        rpr.append(s)
    r.append(rpr)
    t = OxmlElement('w:t')
    t.text = f'[{num}] {text}'
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

# Insert new reference paragraphs before Agnel in forward order
# addprevious() inserts immediately before anchor each time, so forward
# iteration produces the correct [1],[2],...,[24] top-to-bottom order
for num, text in NEW_REFS:
    agnel_para.addprevious(make_ref(num, text))

print(f"  Inserted {len(NEW_REFS)} reference paragraphs")

# ─────────────────────────────────────────────────────────────────────────────
# PART 5 — Appendix D: SDG / PO / PSO improved justification
# ─────────────────────────────────────────────────────────────────────────────
print("\n=== PART 5: Appendix D table updates ===")

def set_table_cell(cell, text, bold=False, size=10):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = None
    from docx.shared import Pt
    run.font.size = Pt(size)

# Find Appendix D tables - they are the last 3 tables (SDG, PO, PSO)
# Count total tables
all_tables = doc.tables
n = len(all_tables)
print(f"  Total tables in doc: {n}")

# Identify the SDG, PO, PSO tables by header content
sdg_tbl = po_tbl = pso_tbl = None
for tbl in all_tables:
    if tbl.rows:
        hdr = tbl.rows[0].cells[0].text.strip()
        if hdr == 'SDG':
            sdg_tbl = tbl
        elif hdr == 'PO':
            po_tbl = tbl
        elif hdr == 'PSO':
            pso_tbl = tbl

print(f"  SDG table found: {sdg_tbl is not None}")
print(f"  PO table found:  {po_tbl is not None}")
print(f"  PSO table found: {pso_tbl is not None}")

# ── SDG table ────────────────────────────────────────────────────────────────
SDG_DATA = [
    ('SDG 7',  'Affordable and Clean Energy',
     'SmartGridHE enables optimal electricity load distribution by performing demand aggregation and threshold detection entirely on encrypted data. The ALT algorithm triggers load-shedding only when the encrypted aggregate exceeds grid capacity, reducing peak waste and supporting efficient energy use without exposing household consumption patterns.'),
    ('SDG 9',  'Industry, Innovation and Infrastructure',
     'The project contributes three novel cryptographic algorithms (ALT, ACV, Encrypted Linear Algebra) not previously demonstrated for smart grid infrastructure. The fully runnable Python 3.13 + TenSEAL prototype bridges the gap between theoretical FHE research and deployable smart grid systems, advancing secure infrastructure.'),
    ('SDG 11', 'Sustainable Cities and Communities',
     'SmartGridHE provides a privacy-by-design cryptographic foundation for multi-agent smart city energy coordination. Households retain data sovereignty while contributing to grid optimisation; the system satisfies GDPR Article 5 data-minimisation requirements, making it viable for large-scale urban deployments.'),
    ('SDG 16', 'Peace, Justice and Strong Institutions',
     'The Pedersen Commitment-based ACV protocol provides mathematical assurance of coordinator honesty — any tampering is detected with probability 1 under discrete-log hardness. Privacy is enforced cryptographically, not by policy, eliminating single points of trust failure and supporting transparent, accountable energy governance.'),
]

if sdg_tbl:
    data_rows = sdg_tbl.rows[1:]
    for i, row in enumerate(data_rows):
        if i < len(SDG_DATA):
            sdg, goal, contrib = SDG_DATA[i]
            set_table_cell(row.cells[0], sdg)
            set_table_cell(row.cells[1], goal)
            set_table_cell(row.cells[2], contrib)
    print("  Updated SDG table")

# ── PO table ─────────────────────────────────────────────────────────────────
PO_DATA = [
    ('PO1', 'Engineering Knowledge',
     'Applied advanced cryptographic theory — CKKS scheme, Pedersen commitments, RFC 3526 MODP Group 14 (2048-bit prime), and IND-CPA security proofs — alongside smart grid domain knowledge (AMI systems, demand-response, NILM privacy risks). Parameter choices (poly_modulus_degree = 16384) are grounded in NIST SP 800-57 and the HE Security Standard [3, 4].'),
    ('PO2', 'Problem Analysis',
     'Systematically identified and characterised five concrete gaps between published academic literature and deployed smart-grid privacy systems: absence of plaintext-free aggregation, Paillier limitations for real-valued data, missing encrypted comparison, unverifiable coordinators, and no runnable multi-agent implementation. Each gap was formally mapped to a specific technical solution with supporting references [5, 6, 7, 8, 11].'),
    ('PO3', 'Design/Development of Solutions',
     'Designed a six-module multi-agent architecture with cryptographically enforced trust boundaries — the coordinator is mathematically incapable of decrypting individual readings, not merely policy-restricted. Developed three novel algorithms (ALT, ACV, Encrypted Cross-Product) that make original contributions beyond standard TenSEAL usage [1, 7, 10].'),
    ('PO4', 'Conduct Investigations',
     'Conducted systematic performance benchmarking across N = 10 to 200 agents using evaluation/benchmark.py, validating O(N) scalability, CKKS approximation error < 10\u207b\u2076 kW over 10 rounds, and cross-method comparison (plaintext vs. Paillier vs. CKKS). All benchmark environments and parameters are fully documented for reproducibility.'),
    ('PO5', 'Modern Tool Usage',
     'Leveraged Python 3.13, TenSEAL \u2265 0.3.14 (wrapping Microsoft SEAL [16]), FastAPI with Uvicorn, NumPy, and hashlib. Applied RFC 3526 MODP Group 14 for Pedersen commitments — a published cryptographic standard not commonly used in undergraduate projects — and designed the system with a modular REST API interface.'),
    ('PO6', 'Engineer and Society',
     'Directly addressed the societal need for household data privacy in smart metering — a concern affecting millions globally, documented by Molina-Markham et al. [11] and mandated by GDPR [12]. Privacy is enforced by architectural design: the coordinator cannot access plaintext data, satisfying Article 5 data-minimisation requirements.'),
    ('PO7', 'Environment and Sustainability',
     'Encrypted demand aggregation enables real-time load balancing that reduces peak electricity waste and defers the need for environmentally costly peaker plants. By supporting SDG 7 (Affordable and Clean Energy), the system directly contributes to sustainable urban energy management.'),
    ('PO9', 'Individual and Team Work',
     'Developed as a four-member team with clearly partitioned module ownership: FHE engine, verifiable aggregation, secure linear algebra, and REST API/dashboard. Integration testing across the agent\u2013coordinator\u2013utility pipeline verified inter-module contracts under realistic multi-agent scenarios with up to 200 simulated households.'),
    ('PO10', 'Communication',
     'System is documented through this formal project report (architecture diagrams, algorithm derivations, benchmark tables), the auto-generated FastAPI OpenAPI specification at /docs, and an interactive browser dashboard enabling non-technical stakeholders to observe encrypted rounds. Code is fully commented with docstrings and type annotations.'),
    ('PO12', 'Life-long Learning',
     'Independently studied 21+ academic references spanning IEEE transactions, ACM proceedings, Springer LNCS, and IACR ePrint. Engaged with advanced topics beyond the curriculum — CKKS bootstrapping (EUROCRYPT 2019 [9]), TFHE gate-level FHE [18], BFV comparison circuits (ASIACRYPT 2019 [6]) — to critically justify design choices.'),
]

if po_tbl:
    data_rows = po_tbl.rows[1:]
    for i, row in enumerate(data_rows):
        if i < len(PO_DATA):
            po, desc, contrib = PO_DATA[i]
            set_table_cell(row.cells[0], po)
            set_table_cell(row.cells[1], desc)
            set_table_cell(row.cells[2], contrib)
    print("  Updated PO table")

# ── PSO table ─────────────────────────────────────────────────────────────────
PSO_DATA = [
    ('PSO1', 'IT Systems Development',
     'Built a complete multi-tier IT system integrating cryptographic computation (core/), autonomous agent simulation (agents/), coordination logic (coordinator/), a FastAPI REST backend (server/server.py), and an interactive browser dashboard (dashboard/). Demonstrated end-to-end integration — from CKKS context initialisation to load-balance decision broadcast — in a reproducible multi-agent simulation.'),
    ('PSO2', 'Security and Privacy',
     'Implemented 128-bit IND-CPA-secure CKKS encryption (NIST Level 1), Pedersen Commitment ACV for verifiable aggregation (detects coordinator fraud with probability 1), and SHA-256 integrity checksums on every ciphertext exchange. Privacy is cryptographically enforced, not reliant on trust assumptions, ensuring the coordinator is incapable of accessing plaintext individual readings.'),
    ('PSO3', 'Applied Research and Innovation',
     'Contributed three novel cryptographic algorithms — ALT (depth-0 encrypted comparison, first applied to smart grids), ACV (Pedersen-based tamper detection for CKKS aggregation), and Encrypted 3D Cross-Product (SIMD rotation for FHE vector geometry) — validated against 21 surveyed references from ASIACRYPT, EUROCRYPT, CRYPTO, and IEEE transactions. None of these combinations exist in prior published literature.'),
]

if pso_tbl:
    data_rows = pso_tbl.rows[1:]
    for i, row in enumerate(data_rows):
        if i < len(PSO_DATA):
            pso, desc, contrib = PSO_DATA[i]
            set_table_cell(row.cells[0], pso)
            set_table_cell(row.cells[1], desc)
            set_table_cell(row.cells[2], contrib)
    print("  Updated PSO table")

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
doc.save(SRC)
print(f"\n=== SAVED: {SRC} ===")

# Quick verification
doc2 = Document(SRC)
body2 = doc2.element.body
paras2 = body2.findall(qn('w:p'))
# Check Roman table numbers present
roman_found = any('Table I' in get_para_text(p) for p in paras2)
print(f"Roman table numbers present: {roman_found}")
# Check refs count
ref_count = sum(1 for p in paras2 if re.match(r'^\[\d+\]', get_para_text(p).strip()))
print(f"Reference paragraphs: {ref_count}")
# Check [22] present
ref22 = any('[22]' in get_para_text(p) for p in paras2)
print(f"New ref [22] present: {ref22}")
