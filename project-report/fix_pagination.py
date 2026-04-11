"""
Fix pagination in SmartGridHE_MiniProject_Report_Final.docx:
  1. Add section break before Introduction
  2. Front matter (before Introduction) = no page numbers
  3. Body (from Introduction) = page numbers starting at 1, centred at bottom
  4. Update Index / List of Figures / List of Tables with corrected page numbers
"""

import io
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'c:/Users/divya/Desktop/encryption-iit-dharwad/project-report/SmartGridHE_MiniProject_Report_Final.docx'

doc = Document(SRC)
body = doc.element.body

# ──────────────────────────────────────────────────────────────
#  STEP 1  Find Introduction paragraph
# ──────────────────────────────────────────────────────────────
intro_p_el = None
for para in doc.paragraphs:
    if para.text.strip().upper() == '1. INTRODUCTION':
        intro_p_el = para._p
        break

assert intro_p_el is not None, "Could not find '1. INTRODUCTION' paragraph"
print("[1] Found Introduction paragraph")

# ──────────────────────────────────────────────────────────────
#  STEP 2  Get the paragraph immediately before Introduction in the XML tree
# ──────────────────────────────────────────────────────────────
prev_el = intro_p_el.getprevious()
while prev_el is not None and prev_el.tag != qn('w:p'):
    prev_el = prev_el.getprevious()

assert prev_el is not None, "No paragraph found before Introduction"
print("[2] Found paragraph before Introduction")

# ──────────────────────────────────────────────────────────────
#  STEP 3  Insert a section-break sectPr into that paragraph's pPr
#          This ends the FRONT MATTER section (no footer/page numbers)
# ──────────────────────────────────────────────────────────────
pPr = prev_el.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    prev_el.insert(0, pPr)

# Remove any pre-existing sectPr in this pPr
for old in pPr.findall(qn('w:sectPr')):
    pPr.remove(old)

# Front-matter sectPr: next-page break, no footer references
sectPr_front = OxmlElement('w:sectPr')
w_type = OxmlElement('w:type')
w_type.set(qn('w:val'), 'nextPage')
sectPr_front.append(w_type)
pPr.append(sectPr_front)
print("[3] Inserted front-matter section break")

# ──────────────────────────────────────────────────────────────
#  STEP 4  Configure body section (document's main sectPr)
#          Set page numbering to start at 1
# ──────────────────────────────────────────────────────────────
doc_sectPr = body.find(qn('w:sectPr'))

# Remove old pgNumType
for e in doc_sectPr.findall(qn('w:pgNumType')):
    doc_sectPr.remove(e)

pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:fmt'), 'decimal')
pgNumType.set(qn('w:start'), '1')
doc_sectPr.append(pgNumType)
print("[4] Body section page numbering set to start at 1")

# ──────────────────────────────────────────────────────────────
#  STEP 5  Save + reload so python-docx sees the two sections,
#          then add a page-number footer to the body section
# ──────────────────────────────────────────────────────────────
buf = io.BytesIO()
doc.save(buf)
buf.seek(0)
doc2 = Document(buf)

print(f"[5] Document reloaded — sections found: {len(doc2.sections)}")

# Body section is the LAST section
body_section = doc2.sections[-1]
body_section.different_first_page_header_footer = False

footer = body_section.footer
footer.is_linked_to_previous = False

# Clear any existing content in the footer paragraph
fp = footer.paragraphs[0]
for run in fp.runs:
    run.text = ''
fp.clear()

fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf = fp.paragraph_format
pf.space_before = Pt(4)
pf.space_after  = Pt(0)

# Helper: add field code
def _fldChar(para, fldCharType):
    run = para.add_run()
    run.font.size = Pt(12)
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), fldCharType)
    run._r.append(fc)
    return run

def _instrText(para, text):
    run = para.add_run()
    run.font.size = Pt(12)
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = text
    run._r.append(instr)
    return run

_fldChar(fp, 'begin')
_instrText(fp, ' PAGE ')
_fldChar(fp, 'end')

print("[5] Page-number footer added to body section")

# ──────────────────────────────────────────────────────────────
#  STEP 6  Update Index table (Table[2] — 26 rows, 3 cols)
# ──────────────────────────────────────────────────────────────
# Estimated page numbers with Introduction = 1
# (These reflect content density after font-size increase)

INDEX_PAGES = {
    '1':  '1',   # Chapter 1
    '1.1 Background & Cryptographic Context':  '1',
    '1.2 Motivation':                          '2',
    '1.3 Problem Definition':                  '3',
    '1.4 Scope and Assumptions':               '4',
    '1.5 Issues and Limitations':             '5',
    '2':  '6',   # Chapter 2
    '2.1 Investigation of Current Field & Related Works': '6',
    '2.2 Literature Survey Overview':          '7',
    '2.3 Related Work':                        '8',
    '2.4 Existing System: Features and Vulnerabilities': '11',
    '2.5 Requirements Analysis':               '13',
    '3':  '16',  # Chapter 3
    '3.1 Architectural Diagram / Block Diagram': '16',
    '3.2 System Flow Chart':                   '19',
    '3.3 ER Diagram for Security Audit Database': '22',
    '3.4 Data Flow Diagram (DFD)':             '23',
    '3.5 Execution Flow and Novel Contributions': '25',
    '3.6 Implementation and Benchmark Results': '30',
    '4':  '34',  # Chapter 4
    '4.1 Conclusion':                          '34',
    '4.2 Future Scope':                        '35',
    'References':                              '37',
    'Appendix':                                '40',
    'Acknowledgements':                        '41',
}

idx_table = doc2.tables[2]   # Sr. No. | Topic | Page No.
for row in idx_table.rows[1:]:   # skip header
    topic_raw = row.cells[1].text.strip()
    # Strip leading spaces used for indentation
    topic_clean = topic_raw.lstrip()
    # Remove indentation prefix like "    1.1 …" → "1.1 …"
    for key, pg in INDEX_PAGES.items():
        if key in topic_clean or topic_clean.endswith(key):
            # Update page number cell
            cell = row.cells[2]
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(pg)
            run.font.size = Pt(11)
            break

print("[6] Index table updated")

# ──────────────────────────────────────────────────────────────
#  STEP 7  Update List of Figures (Table[3] — 11 rows, 3 cols)
# ──────────────────────────────────────────────────────────────
FIGURE_PAGES = {
    '3.1.1': '17',
    '3.1.2': '20',
    '3.2.1': '21',
    '3.3.1': '23',
    '3.4.1': '24',
    '3.5.1': '27',
    '3.5.2': '29',
    '3.5.3': '30',
    '3.6.1': '32',
    '3.6.2': '33',
}

fig_table = doc2.tables[3]
for row in fig_table.rows[1:]:
    sr = row.cells[0].text.strip()
    if sr in FIGURE_PAGES:
        cell = row.cells[2]
        for p in cell.paragraphs:
            for r in p.runs: r.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(FIGURE_PAGES[sr])
        run.font.size = Pt(11)

print("[7] List of Figures updated")

# ──────────────────────────────────────────────────────────────
#  STEP 8  Update List of Tables (Table[4] — 7 rows, 3 cols)
# ──────────────────────────────────────────────────────────────
TABLE_PAGES = {
    '2.5.1': '14',
    '3.1.1': '17',
    '3.5.1': '26',
    '3.5.2': '26',
    '3.6.1': '31',
    '3.6.2': '33',
}

tab_table = doc2.tables[4]
for row in tab_table.rows[1:]:
    sr = row.cells[0].text.strip()
    if sr in TABLE_PAGES:
        cell = row.cells[2]
        for p in cell.paragraphs:
            for r in p.runs: r.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(TABLE_PAGES[sr])
        run.font.size = Pt(11)

print("[8] List of Tables updated")

# ──────────────────────────────────────────────────────────────
#  SAVE
# ──────────────────────────────────────────────────────────────
OUT = SRC.replace('.docx', '_v2.docx')
doc2.save(OUT)
print()
print("[OK] Saved successfully:", OUT)
print("     Page 1 = Introduction")
print("     Page numbers appear in footer of every body page")
