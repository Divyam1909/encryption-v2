"""
update_report.py  —  Apply all edits to SmartGridHE report backup.
"""

import copy
import shutil
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from lxml import etree

# ──────────────────────────────────────────────────────────────────────────────
# Paths
# ──────────────────────────────────────────────────────────────────────────────
FOLDER     = r"c:\Users\divya\Desktop\encryption-iit-dharwad\project-report"
ORIGINAL   = os.path.join(FOLDER, "SmartGridHE_MiniProject_Report_Final_v2.docx")
BACKUP     = os.path.join(FOLDER, "SmartGridHE_MiniProject_Report_Final_v2_BACKUP.docx")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 1 — Create backup
# ══════════════════════════════════════════════════════════════════════════════
shutil.copy2(ORIGINAL, BACKUP)
print(f"[TASK 1] Backup created: {BACKUP}")

doc = Document(BACKUP)

# ──────────────────────────────────────────────────────────────────────────────
# Helper: build a sectPr page-border element
# ──────────────────────────────────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def make_border_element(val="single", sz="18", space="24", color="000000",
                        offset_from="page"):
    """Return a fully populated w:pgBorders element."""
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), offset_from)
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   val)
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), space)
        b.set(qn("w:color"), color)
        pgBorders.append(b)
    return pgBorders


def apply_border_to_sectpr(sectPr):
    """Remove any existing pgBorders, then add a fresh one."""
    for old in sectPr.findall(qn("w:pgBorders")):
        sectPr.remove(old)
    sectPr.append(make_border_element())


# ══════════════════════════════════════════════════════════════════════════════
# TASK 2 — Page borders on ALL sections
# ══════════════════════════════════════════════════════════════════════════════
body = doc.element.body

# Main sectPr (last child of body)
main_sectPr = body.find(qn("w:sectPr"))
if main_sectPr is not None:
    apply_border_to_sectpr(main_sectPr)
    print("[TASK 2] Border applied to main sectPr")

# Inline sectPr elements inside w:pPr
count_inline = 0
for pPr in body.iter(qn("w:pPr")):
    inline_sectPr = pPr.find(qn("w:sectPr"))
    if inline_sectPr is not None:
        apply_border_to_sectpr(inline_sectPr)
        count_inline += 1
print(f"[TASK 2] Border applied to {count_inline} inline sectPr element(s)")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 3 — Update INDEX table
# ══════════════════════════════════════════════════════════════════════════════
# Index mapping: topic text (partial) → correct page number string
INDEX_MAP = {
    "Introduction":                                  "1",
    "1.1":                                           "1",
    "1.2":                                           "2",
    "1.3":                                           "3",
    "1.4":                                           "3",
    "1.5":                                           "4",
    "Literature Survey":                             "5",
    "2.1":                                           "5",
    "2.2":                                           "5",
    "2.3":                                           "6",
    "2.4":                                           "8",
    "2.5":                                           "9",
    "System Design":                                 "11",
    "3.1":                                           "11",
    "3.2":                                           "13",
    "3.3":                                           "15",
    "3.4":                                           "17",
    "3.5":                                           "18",
    "3.6":                                           "24",
    "Conclusion":                                    "28",
    "4.1":                                           "28",
    "4.2":                                           "29",
    "References":                                    "31",
    "Appendix":                                      "35",   # will be renamed
    "Acknowledgements":                              "39",
}

def set_cell_text(cell, text, bold=False, font_size=None):
    """Clear a cell and set its text, preserving paragraph formatting."""
    para = cell.paragraphs[0]
    # clear runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        run = para.runs[0]
        run.text = text
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
    else:
        run = para.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)

def get_cell_text(cell):
    return "".join(p.text for p in cell.paragraphs).strip()


index_table = doc.tables[2]
appendix_row_idx = None

for ri, row in enumerate(index_table.rows):
    topic_text = get_cell_text(row.cells[1])
    page_cell  = row.cells[2]

    # Rename "Appendix" row
    if topic_text.strip() == "Appendix":
        set_cell_text(row.cells[1], "Appendix A: Code Samples")
        set_cell_text(page_cell, "35")
        appendix_row_idx = ri
        print(f"[TASK 3] Renamed 'Appendix' row to 'Appendix A: Code Samples', page=35 (row {ri})")
        continue

    # Generic page updates
    for key, page in INDEX_MAP.items():
        if key == "Appendix":
            continue
        if topic_text.startswith(key) or (key in topic_text and len(key) > 4):
            current = get_cell_text(page_cell)
            if current != page:
                set_cell_text(page_cell, page)
                print(f"[TASK 3] Updated page for '{topic_text[:50]}': {current} -> {page}")
            break

# ── Add new rows after Acknowledgements ──────────────────────────────────────
NEW_ROWS = [
    ("", "Appendix B: Mini Project Report Appendix", "38"),
    ("", "Appendix C: Software Requirement Specification (SRS)", "40"),
    ("", "Appendix D: Mapping to SDGs, POs, and PSOs", "42"),
    ("", "Appendix E: Project Commencement Letters", "43"),
]

# Find the Acknowledgements row
ack_row_elem = None
for row in index_table.rows:
    if "Acknowledgements" in get_cell_text(row.cells[1]):
        ack_row_elem = row._tr
        break

if ack_row_elem is not None:
    # Use a template row (Acknowledgements row) to clone structure
    template_tr = ack_row_elem
    for (sr, topic, page) in reversed(NEW_ROWS):
        new_tr = copy.deepcopy(template_tr)
        cells  = new_tr.findall(qn("w:tc"))
        # Set cell texts
        def _set_tc_text(tc, text):
            for p in tc.findall(qn("w:p")):
                tc.remove(p)
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.text = text
            if " " in text:
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
        _set_tc_text(cells[0], sr)
        _set_tc_text(cells[1], topic)
        _set_tc_text(cells[2], page)
        ack_row_elem.addnext(new_tr)
        print(f"[TASK 3] Inserted new row: '{topic}' p.{page}")
else:
    print("[TASK 3] WARNING: Acknowledgements row not found — new rows not inserted")

print("[TASK 3] Index table update complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 4 — Update LOF and LOT tables
# ══════════════════════════════════════════════════════════════════════════════
LOF_MAP = {
    "3.1.1": "12",
    "3.1.2": "13",
    "3.2.1": "15",
    "3.3.1": "16",
    "3.4.1": "17",
    "3.5.1": "21",
    "3.5.2": "22",
    "3.5.3": "24",
    "3.6.1": "26",
    "3.6.2": "26",
}
LOT_MAP = {
    "2.5.1": "10",
    "3.1.1": "12",
    "3.5.1": "18",
    "3.5.2": "20",
    "3.6.1": "24",
    "3.6.2": "27",
}

def update_page_table(table, sr_map, label):
    for row in table.rows:
        sr   = get_cell_text(row.cells[0]).strip()
        if sr in sr_map:
            old = get_cell_text(row.cells[2])
            new = sr_map[sr]
            if old != new:
                set_cell_text(row.cells[2], new)
                print(f"[TASK 4] {label} {sr}: {old} -> {new}")

# LOF = table 3, LOT = table 4
update_page_table(doc.tables[3], LOF_MAP, "LOF")
update_page_table(doc.tables[4], LOT_MAP, "LOT")
print("[TASK 4] LOF and LOT update complete")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 5 — Insert Appendix A: Code Samples
# ══════════════════════════════════════════════════════════════════════════════

def make_page_break():
    p   = OxmlElement("w:p")
    r   = OxmlElement("w:r")
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def make_para(text, bold=False, italic=False, font_name="Times New Roman",
              font_size=11, align=None, space_preserve=True):
    """Create a w:p element with given formatting."""
    p   = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    if align == "center":
        jc.set(qn("w:val"), "center")
    else:
        jc.set(qn("w:val"), "both")
    pPr.append(jc)
    p.append(pPr)

    if not text:
        return p

    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),   font_name)
    rPr.append(rFonts)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:i")
        rPr.append(i)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    rPr.append(szCs)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    if space_preserve:
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p


def make_code_para(line, font_size=8):
    """Create a single-line code paragraph in Courier New."""
    p   = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)
    # no space before/after
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "0")
    spacing.set(qn("w:line"),   "240")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    p.append(pPr)

    if line == "":
        return p

    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),  "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.append(rFonts)
    sz   = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = line
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    return p


CODE_A1 = """\
import tenseal as ts, hashlib
from datetime import datetime

class SmartGridFHE:
    \"\"\"CKKS FHE Engine \u2014 Coordinator gets PUBLIC context (no decryption);
       only Utility Company holds SECRET context (128-bit NIST Level 1).\"\"\"

    def __init__(self, poly_modulus_degree=16384,
                 coeff_mod_bit_sizes=None, global_scale=2**40):
        if coeff_mod_bit_sizes is None:
            coeff_mod_bit_sizes = [60, 40, 40, 40, 60]  # 3-4 mult. depths
        self.context = ts.context(
            ts.SCHEME_TYPE.CKKS,
            poly_modulus_degree=poly_modulus_degree,
            coeff_mod_bit_sizes=coeff_mod_bit_sizes)
        self.context.generate_galois_keys()   # SIMD cyclic rotation
        self.context.generate_relin_keys()    # Post-multiply relinearisation
        self.context.global_scale  = global_scale
        self.context.auto_rescale  = True
        self.context.auto_relin    = True
        self.context.auto_mod_switch = True

    def get_public_context(self) -> bytes:
        \"\"\"Serialise context WITHOUT secret key \u2014 distribute to agents/coordinator.\"\"\"
        ctx = self.context.copy()
        ctx.make_context_public()
        return ctx.serialize()

    def is_private(self) -> bool:
        return self.context.is_private()

    def encrypt_demand(self, demand_kw: float, agent_id: str):
        \"\"\"Real-valued CKKS encryption of demand d_i \u2192 E(d_i), ~42 KB ciphertext.\"\"\"
        encrypted  = ts.ckks_vector(self.context, [float(demand_kw)])
        ciphertext = encrypted.serialize()
        checksum   = hashlib.sha256(ciphertext).hexdigest()[:12]
        return EncryptedDemand(ciphertext, datetime.now().isoformat(),
                               agent_id, 1, checksum)

    def aggregate_demands(self, encrypted_list: list):
        \"\"\"E(\u03a3d_i) = \u03a3E(d_i)  [ciphertext addition, depth 0].\"\"\"
        result = ts.ckks_vector_from(self.context, encrypted_list[0].ciphertext)
        for enc in encrypted_list[1:]:
            result = result + ts.ckks_vector_from(self.context, enc.ciphertext)
        return self._save(result, encrypted_list[0], "aggregate")

    def compute_average(self, encrypted_total, count: int):
        \"\"\"E(avg) = E(\u03a3d) \u00d7 (1/N)  [scalar multiply, depth 0].\"\"\"
        vec    = ts.ckks_vector_from(self.context, encrypted_total.ciphertext)
        result = vec * (1.0 / count)
        return self._save(result, encrypted_total, "average")

    def decrypt_demand(self, encrypted) -> list:
        \"\"\"Decrypt aggregate \u2014 ONLY callable with secret context (Utility only).\"\"\"
        if not self.context.is_private():
            raise ValueError("Cannot decrypt: no secret key.")
        return ts.ckks_vector_from(self.context, encrypted.ciphertext).decrypt()\
"""

CODE_A2 = """\
class EncryptedThresholdDetector:
    \"\"\"Encrypted comparison at multiplicative depth 0.
    Prior polynomial methods need depth 8-15 [4,5]; ALT uses none.\"\"\"

    def __init__(self, fhe_engine, default_sensitivity=7.0):
        self.fhe = fhe_engine

    def detect_threshold_encrypted(self, encrypted_value,
                                   threshold: float, sensitivity: float = 7.0):
        \"\"\"
        ALT score:  s(x) = 0.5 + (x - T) \u00d7 (0.5/\u03b4),  \u03b4 = T/k
        Rearranged: E(s) = E(x) \u00d7 slope + intercept       [depth 0]
          slope     = 0.5k / T
          intercept = 0.5 - 0.5k  (may be negative; clamped post-decryption)
        Properties: s(T) = 0.5 (uncertain), s(T-\u03b4) = 0.0, s(T+\u03b4) = 1.0
        \"\"\"
        k         = sensitivity
        delta     = threshold / k
        slope     = 0.5 / delta              # = 0.5k / T
        intercept = 0.5 - threshold * slope  # = 0.5 - 0.5k

        vec_x  = ts.ckks_vector_from(self.fhe.context, encrypted_value.ciphertext)
        result = vec_x * slope + intercept   # multiply_plain then add_plain

        ct = result.serialize()
        return EncryptedDemand(ct, datetime.now().isoformat(), "alt_score",
                               1, hashlib.sha256(ct).hexdigest()[:12])

    @staticmethod
    def interpret_score(score: float) -> tuple:
        \"\"\"Utility interprets decrypted score. Returns (zone, action).\"\"\"
        s = max(0.0, min(1.0, score))        # clamp to [0, 1]
        if   s < 0.3: return 'below',     'NORMAL'
        elif s > 0.7: return 'above',     'REDUCE / CRITICAL'
        else:         return 'uncertain', 'MONITOR'\
"""

CODE_A3 = """\
import secrets, hashlib

# RFC 3526 MODP Group 14 (2048-bit prime)  \u2014  truncated for brevity
PEDERSEN_PRIME = int("FFFFFFFF...FFFFFFFF", 16)   # full value in source
PEDERSEN_G     = 2
PEDERSEN_H     = pow(PEDERSEN_G,
                     int.from_bytes(hashlib.sha256(
                         b"SmartGridHE_Pedersen_H_v1").digest(), 'big'),
                     PEDERSEN_PRIME)

class PedersenCommitmentScheme:
    \"\"\"C_i = g^(d_i\u00b7s) \u00b7 h^r_i mod p  (hiding + binding + additive hom.)
    C(a) \u00d7 C(b) = C(a+b)  enables aggregation verification.\"\"\"

    SCALE = 1_000_000     # float \u2192 integer scaling (6 decimal places)

    def commit(self, value: float) -> tuple:
        \"\"\"Returns (C_public \u2192 coordinator,  r_secret \u2192 utility via secure ch.)\"\"\"
        m = int(value * self.SCALE)
        r = secrets.randbelow(PEDERSEN_PRIME - 1)
        C = (pow(PEDERSEN_G, m, PEDERSEN_PRIME) *
             pow(PEDERSEN_H, r, PEDERSEN_PRIME)) % PEDERSEN_PRIME
        return C, r

    def aggregate_commitments(self, commitments: list) -> int:
        \"\"\"C_agg = \u220fC_i mod p = g^(\u03a3d_i\u00b7s) \u00b7 h^(\u03a3r_i) mod p\"\"\"
        agg = 1
        for C in commitments:
            agg = (agg * C) % PEDERSEN_PRIME
        return agg

    def verify(self, claimed_sum: float, C_agg: int, r_total: int) -> bool:
        \"\"\"
        Utility checks: C_agg == g^(\u03a3d\u00b7s) \u00b7 h^(\u03a3r) mod p
        True  \u2192 VALID (coordinator computed correctly)
        False \u2192 COORDINATOR CHEATED (detected with probability 1)
        \"\"\"
        m_total  = int(claimed_sum * self.SCALE)
        expected = (pow(PEDERSEN_G, m_total,  PEDERSEN_PRIME) *
                    pow(PEDERSEN_H, r_total % (PEDERSEN_PRIME-1),
                        PEDERSEN_PRIME)) % PEDERSEN_PRIME
        return expected == C_agg\
"""

# Find the target paragraph (first para containing "Agnel Charities")
target_para = None
for para in doc.paragraphs:
    if "Agnel Charities" in para.text:
        target_para = para._element
        break

if target_para is None:
    print("[TASK 5] WARNING: 'Agnel Charities' paragraph not found — Appendix A not inserted")
else:
    # Build elements to insert (will be inserted in reversed order with addprevious)
    elements_to_insert = []

    # 1. Page break
    elements_to_insert.append(make_page_break())

    # 2. Heading
    elements_to_insert.append(
        make_para("APPENDIX A: CODE SAMPLES", bold=True, align="center",
                  font_name="Times New Roman", font_size=14))

    # 3. Blank line
    elements_to_insert.append(make_para(""))

    # 4. Intro paragraph
    elements_to_insert.append(
        make_para(
            "This appendix presents condensed implementations of the three core modules "
            "of SmartGridHE. Full source code is available in the smart-grid-he/ directory "
            "of the project repository.",
            align="both", font_name="Times New Roman", font_size=11))

    # 5. Blank line
    elements_to_insert.append(make_para(""))

    # 6. Sub-heading A.1
    elements_to_insert.append(
        make_para("A.1  FHE Engine \u2013 CKKS Context Setup and Core Operations  (core/fhe_engine.py)",
                  bold=True, font_name="Times New Roman", font_size=11, align="left"))

    # 7. Code block A.1
    for line in CODE_A1.split("\n"):
        elements_to_insert.append(make_code_para(line))

    # 8. Blank line
    elements_to_insert.append(make_para(""))

    # 9. Sub-heading A.2
    elements_to_insert.append(
        make_para("A.2  Novel Contribution #1 \u2013 Adaptive Linear Threshold (ALT)  (core/polynomial_comparator.py)",
                  bold=True, font_name="Times New Roman", font_size=11, align="left"))

    # 10. Code block A.2
    for line in CODE_A2.split("\n"):
        elements_to_insert.append(make_code_para(line))

    # 11. Blank line
    elements_to_insert.append(make_para(""))

    # 12. Sub-heading A.3
    elements_to_insert.append(
        make_para("A.3  Novel Contribution #2 \u2013 Pedersen Commitment Verifiable Aggregation  (core/verifiable_aggregation.py)",
                  bold=True, font_name="Times New Roman", font_size=11, align="left"))

    # 13. Code block A.3
    for line in CODE_A3.split("\n"):
        elements_to_insert.append(make_code_para(line))

    # 14. Blank line
    elements_to_insert.append(make_para(""))

    # Insert all in REVERSED order using addprevious
    for elem in reversed(elements_to_insert):
        target_para.addprevious(elem)

    print(f"[TASK 5] Appendix A inserted ({len(elements_to_insert)} elements) before 'Agnel Charities'")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 6 — Fill Appendix B table (Group Details)
# ══════════════════════════════════════════════════════════════════════════════
GROUP_LINES = [
    "Chacko Martin Koyikkattuchira (5023124)",
    "Divyam Navin (5023134)",
    "Atharva Palve (5023136)",
    "Akshat Santosh Sawant (5023151)",
    "Guide: Dr. Vaishali Bodade, Ms. Sharlene Rebeiro",
]

def fill_cell_multiline(cell, lines, font_name="Times New Roman", font_size=11):
    """Fill a table cell with multiple paragraphs (one per line)."""
    # Remove existing paragraphs content
    tc = cell._tc
    for p in list(tc.findall(qn("w:p"))):
        tc.remove(p)
    # Add new paragraphs
    for line in lines:
        p   = OxmlElement("w:p")
        r   = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"),  font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)
        sz   = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(int(font_size * 2)))
        rPr.append(sz)
        rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = line
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)

appendix_b_table = doc.tables[11]
for row in appendix_b_table.rows:
    label = get_cell_text(row.cells[0])
    if "Group Details" in label:
        fill_cell_multiline(row.cells[1], GROUP_LINES)
        print("[TASK 6] Group Details cell filled")
        break

# ══════════════════════════════════════════════════════════════════════════════
# TASK 7 — Fix Appendix C formatting + Python version
# ══════════════════════════════════════════════════════════════════════════════
in_c = False
c_count = 0
for para in doc.paragraphs:
    if "Appendix C: Software Requirement" in para.text:
        in_c = True
    if "APPENDIX D" in para.text.upper():
        in_c = False

    if in_c:
        style_name = para.style.name if para.style else ""
        # Justify body text paragraphs
        if style_name in ("Normal", "Body Text", "") or style_name.startswith("Normal"):
            if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                c_count += 1
        # Replace Python 3.x with Python 3.13
        for run in para.runs:
            if "Python 3.x" in run.text:
                run.text = run.text.replace("Python 3.x", "Python 3.13")
                print("[TASK 7] Replaced 'Python 3.x' with 'Python 3.13'")

print(f"[TASK 7] Justified {c_count} paragraph(s) in Appendix C")

# ══════════════════════════════════════════════════════════════════════════════
# TASK 8 — Fix Appendix D formatting + add SDG 11
# ══════════════════════════════════════════════════════════════════════════════
in_d = False
d_count = 0
sdg9_para_elem = None
sdg11_present  = False

for para in doc.paragraphs:
    if "APPENDIX D" in para.text.upper():
        in_d = True
    if "APPENDIX E" in para.text.upper():
        in_d = False

    if in_d:
        style_name = para.style.name if para.style else ""
        if style_name in ("Normal", "Body Text", "") or style_name.startswith("Normal"):
            if para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                d_count += 1

        # Track SDG 9 and SDG 11
        if "SDG 9" in para.text:
            sdg9_para_elem = para._element
        if "SDG 11" in para.text:
            sdg11_present = True

print(f"[TASK 8] Justified {d_count} paragraph(s) in Appendix D")

if not sdg11_present and sdg9_para_elem is not None:
    sdg11_text = (
        "SDG 11 (Sustainable Cities and Communities): The framework supports smart city "
        "energy management through privacy-preserving load balancing at urban scale."
    )
    sdg11_para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc  = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    sdg11_para.append(pPr)

    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b   = OxmlElement("w:b")  # bold for the SDG label
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),  "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "SDG 11 (Sustainable Cities and Communities): "
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    sdg11_para.append(r)

    r2   = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:ascii"),  "Times New Roman")
    rFonts2.set(qn("w:hAnsi"), "Times New Roman")
    rPr2.append(rFonts2)
    r2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = ("The framework supports smart city energy management through "
               "privacy-preserving load balancing at urban scale.")
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r2.append(t2)
    sdg11_para.append(r2)

    sdg9_para_elem.addnext(sdg11_para)
    print("[TASK 8] SDG 11 paragraph inserted after SDG 9")
elif sdg11_present:
    print("[TASK 8] SDG 11 already present — no insertion needed")
else:
    print("[TASK 8] WARNING: SDG 9 paragraph not found — SDG 11 not inserted")

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
doc.save(BACKUP)
print(f"\n[DONE] Backup saved: {BACKUP}")

orig_size   = os.path.getsize(ORIGINAL)
backup_size = os.path.getsize(BACKUP)
print(f"       Original size : {orig_size:,} bytes")
print(f"       Backup size   : {backup_size:,} bytes")
print(f"       Difference    : {backup_size - orig_size:+,} bytes")
if backup_size != orig_size:
    print("       [OK] Files differ — changes confirmed")
else:
    print("       [NOTE] Sizes identical — verify changes were applied")
